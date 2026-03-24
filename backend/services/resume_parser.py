"""Parse PDF or DOCX resume into plain text."""
from pathlib import Path


def parse_resume(file_path: str) -> str:
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"Resume not found: {file_path}")
    ext = path.suffix.lower()
    if ext == ".pdf":
        return _parse_pdf(str(path))
    elif ext in (".docx", ".doc"):
        return _parse_docx(str(path))
    raise ValueError(f"Unsupported resume format: {ext}. Use PDF or DOCX.")


def _parse_pdf(path: str) -> str:
    import pdfplumber
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text(x_tolerance=2, y_tolerance=2)
            if text:
                pages.append(text.strip())
    return "\n\n".join(pages)


def _parse_docx(path: str) -> str:
    from docx import Document
    doc = Document(path)
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    # Also extract text from tables (skills tables are common in DS resumes)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)
